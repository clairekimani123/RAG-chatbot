"""
Updated document_service.py — Week 4 adds support for:
  - Images (PNG, JPEG) — uses Groq vision to describe then embed
  - Plain text files (.txt)
  - Word documents (.docx)

The pipeline stays the same after extraction — only Step 1 changes per type.
"""

import fitz  # PyMuPDF — PDF extraction
from langchain_text_splitters import RecursiveCharacterTextSplitter
import hashlib
import numpy as np
from sqlalchemy.orm import Session
import base64

from app.core.config import get_settings
from app.db.models import Document, Chunk

settings = get_settings()


SUPPORTED_TYPES = {
    "pdf": [".pdf"],
    "image": [".png", ".jpg", ".jpeg", ".webp"],
    "text": [".txt"],
    "docx": [".docx"],
}


def get_file_type(filename: str) -> str:
    """Determine file type from extension."""
    ext = "." + filename.rsplit(".", 1)[-1].lower()
    for ftype, exts in SUPPORTED_TYPES.items():
        if ext in exts:
            return ftype
    raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, PNG, JPEG, TXT, DOCX")


# ── Extractors ────────────────────────────────────────────────────────────────

def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    if not text.strip():
        raise ValueError("No text found in PDF. File may be a scanned image.")
    return text


def extract_from_image(file_bytes: bytes, filename: str) -> str:
    """
    Extract text/description from an image using Groq's vision model.
    Converts image to base64 and sends to llama-3.2-11b-vision-preview.
    """
    from groq import Groq
    client = Groq(api_key=settings.groq_api_key)

    ext = filename.rsplit(".", 1)[-1].lower()
    media_type = "image/jpeg" if ext in ["jpg", "jpeg"] else f"image/{ext}"

    b64 = base64.standard_b64encode(file_bytes).decode("utf-8")

    response = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{media_type};base64,{b64}"},
                },
                {
                    "type": "text",
                    "text": (
                        "Describe this image in detail. Include: "
                        "what you see, any text present, colours, objects, people, "
                        "context, and anything else that would help answer questions about it."
                    ),
                },
            ],
        }],
        max_tokens=1024,
    )
    return response.choices[0].message.content


def extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain .txt file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word .docx file using python-docx."""
    import io
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])


# ── Chunking + Embedding (same for all file types) ───────────────────────────

def chunk_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [c.strip() for c in chunks if c.strip()]

def embed_texts(texts: list[str]) -> list[list[float]]:
    """
    Lightweight deterministic embeddings using hashing.
    No heavy ML models — works within Render's 512MB free tier.
    Uses 384 dimensions to match existing pgvector column size.
    """
    embeddings = []
    for text in texts:
        # Create a 384-dimensional vector from the text
        vector = []
        for i in range(384):
            # Hash the text with different seeds to get different dimensions
            seed = f"{i}:{text[:200]}"
            hash_val = int(hashlib.sha256(seed.encode()).hexdigest(), 16)
            # Normalize to -1 to 1 range
            normalized = (hash_val % 10000) / 5000.0 - 1.0
            vector.append(normalized)
        # Normalize the vector to unit length for cosine similarity
        arr = np.array(vector, dtype=np.float32)
        norm = np.linalg.norm(arr)
        if norm > 0:
            arr = arr / norm
        embeddings.append(arr.tolist())
    return embeddings


def embed_single(text: str) -> list[float]:
    """Embed one string."""
    return embed_texts([text])[0]


# ── Main ingestion pipeline ───────────────────────────────────────────────────

def ingest_document(
    file_bytes: bytes,
    filename: str,
    db: Session,
    user_id: int | None = None,
) -> Document:
    """
    Universal ingestion pipeline — works for all file types.
    Step 1 (extraction) varies by type.
    Steps 2-5 (chunk, embed, save) are identical.
    """
    file_type = get_file_type(filename)

    # Step 1 — extract text based on file type
    if file_type == "pdf":
        raw_text = extract_from_pdf(file_bytes)
    elif file_type == "image":
        raw_text = extract_from_image(file_bytes, filename)
    elif file_type == "text":
        raw_text = extract_from_txt(file_bytes)
    elif file_type == "docx":
        raw_text = extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported type: {file_type}")

    # Steps 2-5 — same for all types
    chunks = chunk_text(raw_text)
    if not chunks:
        raise ValueError("No content found in file.")

    embeddings = embed_texts(chunks)

    document = Document(
        filename=filename,
        original_name=filename,
        file_type=file_type,
        total_chunks=len(chunks),
        user_id=user_id,
    )
    db.add(document)
    db.flush()

    for i, (text, embedding) in enumerate(zip(chunks, embeddings)):
        chunk = Chunk(
            document_id=document.id,
            content=text,
            chunk_index=i,
            embedding=embedding,
        )
        db.add(chunk)

    db.commit()
    db.refresh(document)
    return document


# ── Retrieval ─────────────────────────────────────────────────────────────────

def retrieve_similar_chunks(question: str, document_id: int, db: Session) -> list[Chunk]:
    question_embedding = embed_single(question)
    return (
        db.query(Chunk)
        .filter(Chunk.document_id == document_id)
        .order_by(Chunk.embedding.cosine_distance(question_embedding))
        .limit(settings.top_k_chunks)
        .all()
    )


# ── Document helpers ──────────────────────────────────────────────────────────

def list_documents(db: Session, user_id: int | None = None) -> list[Document]:
    q = db.query(Document)
    if user_id:
        q = q.filter(Document.user_id == user_id)
    return q.order_by(Document.created_at.desc()).all()


def get_document(document_id: int, db: Session) -> Document | None:
    return db.query(Document).filter(Document.id == document_id).first()


def delete_document(document_id: int, db: Session) -> bool:
    doc = get_document(document_id, db)
    if not doc:
        return False
    db.delete(doc)
    db.commit()
    return True